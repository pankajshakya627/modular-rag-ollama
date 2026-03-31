"""Document processing using LangChain text splitters and loaders for Modular RAG.

Uses LangChain built-in text splitters:
- RecursiveCharacterTextSplitter
- CharacterTextSplitter
- langchain_experimental.SemanticChunker

Uses LangChain document loaders:
- PyPDFLoader, Docx2txtLoader, TextLoader

All domain models (DocumentChunk, Document) use Pydantic BaseModel
for runtime validation, serialization, and OpenAPI schema export.
"""
import os
import re
import hashlib
import unicodedata
from enum import Enum
from pathlib import Path
from typing import Any, Dict, List, Optional
import logging

from pydantic import BaseModel, ConfigDict, Field, model_validator

# LangChain text splitters
from langchain_text_splitters import RecursiveCharacterTextSplitter, CharacterTextSplitter
from langchain_core.documents import Document as LangChainDocument

# Direct import to avoid circular dependency
from src.core.uuid_utils import generate_uuid  # UUID v7 for time-sorted IDs

# Import metadata extractor — use try/except to handle circular import gracefully
try:
    from src.utils.metadata_extractor import DocumentMetadataExtractor
    METADATA_EXTRACTION_AVAILABLE = True
except ImportError:
    METADATA_EXTRACTION_AVAILABLE = False

logger = logging.getLogger(__name__)


class ChunkingStrategy(Enum):
    """Available chunking strategies."""
    RECURSIVE = "recursive"
    SEMANTIC = "semantic"
    FIXED_SIZE = "fixed_size"
    SENTENCE = "sentence"
    PARAGRAPH = "paragraph"


class DocumentChunk(BaseModel):
    """Represents a chunk of a document.

    Uses Pydantic for runtime validation, serialization, and schema export.
    """
    model_config = ConfigDict(arbitrary_types_allowed=True)

    id: str = Field(default_factory=lambda: generate_uuid())
    content: str
    metadata: Dict[str, Any] = Field(default_factory=dict)
    embedding: Optional[List[float]] = None
    start_char_idx: int = Field(default=0, ge=0)
    end_char_idx: int = Field(default=0, ge=0)
    chunk_index: int = Field(default=0, ge=0)

    def to_dict(self) -> Dict[str, Any]:
        """Convert chunk to dictionary (excludes embedding for brevity)."""
        data = self.model_dump(exclude={"embedding"})
        return data

    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> "DocumentChunk":
        """Create chunk from dictionary."""
        return cls(**data)


class Document(BaseModel):
    """Represents a processed document.

    Uses Pydantic for runtime validation, serialization, and schema export.
    """
    model_config = ConfigDict(arbitrary_types_allowed=True)

    id: str = Field(default_factory=lambda: generate_uuid())
    content: str
    chunks: List[DocumentChunk] = Field(default_factory=list)
    metadata: Dict[str, Any] = Field(default_factory=dict)
    source_type: str = "text"
    source_path: Optional[str] = None

    def to_dict(self) -> Dict[str, Any]:
        """Convert document to dictionary."""
        return self.model_dump(exclude={"chunks": {"__all__": {"embedding"}}})

    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> "Document":
        """Create document from dictionary."""
        return cls(**data)


def _langchain_docs_to_chunks(
    lc_docs: List[LangChainDocument],
    metadata: Optional[Dict[str, Any]] = None,
) -> List[DocumentChunk]:
    """Convert LangChain Document list to DocumentChunk list."""
    chunks = []
    for idx, lc_doc in enumerate(lc_docs):
        merged_meta = {**(metadata or {}), **lc_doc.metadata}
        chunk = DocumentChunk(
            id=generate_uuid(),
            content=lc_doc.page_content,
            metadata=merged_meta,
            chunk_index=idx,
        )
        chunks.append(chunk)
    return chunks


class DocumentProcessor:
    """Main document processor using LangChain text splitters and loaders."""

    def __init__(
        self,
        chunking_strategy: ChunkingStrategy = ChunkingStrategy.RECURSIVE,
        chunk_size: int = 1024,
        chunk_overlap: int = 128,
        embedding_wrapper=None,
    ):
        """Initialize the document processor."""
        self.chunking_strategy = chunking_strategy
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.embedding_wrapper = embedding_wrapper

        # Metadata extraction settings
        self.enable_metadata_extraction = True
        self.enable_llm_metadata = False
        self.metadata_extractor = None

        self._init_splitter()
        self._init_metadata_extractor()

    def _init_metadata_extractor(self):
        """Initialize metadata extractor if available."""
        if METADATA_EXTRACTION_AVAILABLE and self.enable_metadata_extraction:
            self.metadata_extractor = DocumentMetadataExtractor(
                llm_wrapper=None,
                enable_llm_extraction=self.enable_llm_metadata,
                enable_rule_extraction=True,
            )
            logger.info("Metadata extraction enabled")
        else:
            logger.info("Metadata extraction not available or disabled")

    def _init_splitter(self):
        """Initialize the appropriate LangChain text splitter."""
        if self.chunking_strategy == ChunkingStrategy.RECURSIVE:
            self.splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
                separators=["\n\n", "\n", ". ", "! ", "? ", "; ", " ", ""],
                length_function=len,
            )
        elif self.chunking_strategy == ChunkingStrategy.SEMANTIC:
            if self.embedding_wrapper is None:
                raise ValueError("Embedding wrapper required for semantic chunking")
            try:
                from langchain_experimental.text_splitter import SemanticChunker
                self.splitter = SemanticChunker(
                    embeddings=self.embedding_wrapper,
                    breakpoint_threshold_type="percentile",
                )
            except ImportError:
                logger.warning("langchain-experimental not installed, falling back to recursive")
                self.splitter = RecursiveCharacterTextSplitter(
                    chunk_size=self.chunk_size,
                    chunk_overlap=self.chunk_overlap,
                )
        elif self.chunking_strategy == ChunkingStrategy.FIXED_SIZE:
            self.splitter = CharacterTextSplitter(
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
                separator="",
            )
        elif self.chunking_strategy == ChunkingStrategy.SENTENCE:
            self.splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
                separators=[". ", "! ", "? ", "\n\n", "\n", " ", ""],
            )
        else:
            self.splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
            )

    def process_text(
        self,
        text: str,
        metadata: Optional[Dict[str, Any]] = None,
        doc_id: Optional[str] = None,
    ) -> Document:
        """Process plain text into a document with chunks."""
        doc_id = doc_id or generate_uuid()

        # Clean the text (Unicode-safe, preserves multilingual content)
        text = self._clean_text(text)

        lc_docs = self.splitter.create_documents([text], metadatas=[metadata or {}])
        chunks = _langchain_docs_to_chunks(lc_docs, metadata)

        # Enrich chunks with extracted metadata
        if self.metadata_extractor and chunks:
            chunks = self._enrich_chunks_with_metadata(chunks)

        document = Document(
            id=doc_id,
            content=text,
            chunks=chunks,
            metadata=metadata or {},
            source_type="text",
        )

        logger.info(
            f"Processed text into {len(chunks)} chunks "
            f"(metadata extraction: {self.metadata_extractor is not None})"
        )
        return document

    def _enrich_chunks_with_metadata(self, chunks: List[DocumentChunk]) -> List[DocumentChunk]:
        """Enrich chunks with extracted metadata."""
        for chunk in chunks:
            try:
                extracted = self.metadata_extractor.extract(
                    chunk.content,
                    use_llm=self.enable_llm_metadata,
                )
                extracted_dict = extracted.to_dict()
                chunk.metadata.update(extracted_dict)
            except Exception as e:
                logger.warning(f"Failed to extract metadata for chunk {chunk.id}: {e}")
        return chunks

    def process_file(
        self,
        file_path: str,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> Document:
        """Process a file using LangChain document loaders."""
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        content = self._extract_text_from_file(file_path)

        doc_id = self._generate_file_hash(file_path)

        if metadata is None:
            metadata = {}
        metadata["file_name"] = path.name
        metadata["file_size"] = path.stat().st_size
        metadata["file_type"] = path.suffix

        document = self.process_text(
            text=content,
            metadata=metadata,
            doc_id=doc_id,
        )
        document.source_path = str(path.absolute())
        document.source_type = path.suffix

        return document

    def process_directory(
        self,
        directory_path: str,
        extensions: Optional[List[str]] = None,
        recursive: bool = True,
        max_files: Optional[int] = None,
    ) -> List[Document]:
        """Process all files in a directory."""
        path = Path(directory_path)

        if not path.exists() or not path.is_dir():
            raise ValueError(f"Invalid directory: {directory_path}")

        extensions = extensions or [".txt", ".pdf", ".docx", ".md"]

        files = []
        if recursive:
            for ext in extensions:
                files.extend(path.rglob(f"*{ext}"))
        else:
            for ext in extensions:
                files.extend(path.glob(f"*{ext}"))

        if max_files:
            files = files[:max_files]

        documents = []
        for file_path in files:
            try:
                doc = self.process_file(str(file_path))
                documents.append(doc)
                logger.info(f"Processed: {file_path}")
            except Exception as e:
                logger.error(f"Error processing {file_path}: {e}")

        logger.info(f"Processed {len(documents)} documents from {directory_path}")
        return documents

    def _clean_text(self, text: str) -> str:
        """Clean and normalize text.

        Uses Unicode-safe normalization (NFC) instead of ASCII stripping
        to preserve multilingual content (Chinese, Arabic, etc.).
        """
        # Normalize Unicode to NFC form
        text = unicodedata.normalize("NFC", text)
        # Collapse excessive whitespace
        text = re.sub(r"\s+", " ", text)
        # Remove control characters (but keep printable Unicode)
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
        return text.strip()

    def _extract_text_from_file(self, file_path: str) -> str:
        """Extract text using LangChain document loaders where available."""
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext in (".txt", ".md"):
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()

        elif ext == ".pdf":
            try:
                from langchain_community.document_loaders import PyPDFLoader
                loader = PyPDFLoader(file_path)
                docs = loader.load()
                return "\n".join([doc.page_content for doc in docs])
            except ImportError:
                import pypdf
                text = []
                with open(file_path, "rb") as f:
                    pdf = pypdf.PdfReader(f)
                    for page in pdf.pages:
                        text.append(page.extract_text() or "")
                return "\n".join(text)

        elif ext == ".docx":
            try:
                from langchain_community.document_loaders import Docx2txtLoader
                loader = Docx2txtLoader(file_path)
                docs = loader.load()
                return "\n".join([doc.page_content for doc in docs])
            except ImportError:
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                return "\n".join([para.text for para in doc.paragraphs])

        elif ext == ".html":
            from bs4 import BeautifulSoup
            with open(file_path, "r", encoding="utf-8") as f:
                soup = BeautifulSoup(f.read(), "html.parser")
                return soup.get_text()

        else:
            raise ValueError(f"Unsupported file format: {ext}")

    def _generate_file_hash(self, file_path: str) -> str:
        """Generate a deterministic ID from file content hash."""
        with open(file_path, "rb") as f:
            content = f.read()
            return hashlib.md5(content).hexdigest()

    def set_chunking_strategy(
        self,
        strategy: ChunkingStrategy,
        chunk_size: Optional[int] = None,
        chunk_overlap: Optional[int] = None,
    ):
        """Change the chunking strategy at runtime."""
        self.chunking_strategy = strategy
        if chunk_size:
            self.chunk_size = chunk_size
        if chunk_overlap:
            self.chunk_overlap = chunk_overlap
        self._init_splitter()


# ============================================================================
# Backward compatibility wrappers
# ============================================================================

class RecursiveChunker:
    """Backward-compatible wrapper around LangChain RecursiveCharacterTextSplitter."""

    def __init__(self, chunk_size: int = 1024, chunk_overlap: int = 128, **kwargs):
        self._splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ". ", "! ", "? ", "; ", " ", ""],
        )

    def chunk(self, text: str, metadata: Optional[Dict[str, Any]] = None) -> List[DocumentChunk]:
        if not text:
            return []
        lc_docs = self._splitter.create_documents([text], metadatas=[metadata or {}])
        return _langchain_docs_to_chunks(lc_docs, metadata)


class FixedSizeChunker:
    """Backward-compatible wrapper around LangChain CharacterTextSplitter."""

    def __init__(self, chunk_size: int = 1024, chunk_overlap: int = 128, **kwargs):
        self._splitter = CharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separator="",
        )

    def chunk(self, text: str, metadata: Optional[Dict[str, Any]] = None) -> List[DocumentChunk]:
        if not text:
            return []
        lc_docs = self._splitter.create_documents([text], metadatas=[metadata or {}])
        return _langchain_docs_to_chunks(lc_docs, metadata)
